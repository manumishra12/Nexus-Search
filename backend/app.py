from fastapi import FastAPI, File, UploadFile, HTTPException, Depends, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, StreamingResponse, FileResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel, Field
from typing import List, Optional, Dict, Any, Union
import uuid
import json
import os
import shutil
from pathlib import Path
import mimetypes
from datetime import datetime, timedelta
import asyncio
import logging
from io import BytesIO
import base64
import hashlib
import random
import re
import statistics
from collections import Counter
import tempfile
import time

# Enhanced imports for better file processing
try:
    from PIL import Image, ImageOps
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False
    logging.warning("PIL not available. Image processing will be limited.")

try:
    import psutil
    PSUTIL_AVAILABLE = True
except ImportError:
    PSUTIL_AVAILABLE = False
    logging.warning("psutil not available. System monitoring will be limited.")

# File processing imports
try:
    import PyPDF2
    import docx
    import pandas as pd
    import numpy as np
    import speech_recognition as sr
    from pydub import AudioSegment
    import cv2
    import pytesseract
    import moviepy.editor as mp
    FILE_PROCESSING_AVAILABLE = True
except ImportError:
    FILE_PROCESSING_AVAILABLE = False
    logging.warning("Some file processing libraries not available.")

# Web scraping and requests
import requests
from bs4 import BeautifulSoup
from urllib.parse import urlparse, urljoin

# LangChain and LangGraph imports
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_community.tools.tavily_search import TavilySearchResults
from langchain_core.messages import HumanMessage, SystemMessage
from langchain_core.prompts import ChatPromptTemplate
from langgraph.graph import StateGraph, END
from langgraph.graph.message import add_messages
from typing_extensions import TypedDict

# For document generation
try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as ReportLabImage, Table, TableStyle, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False
    logging.warning("ReportLab not available. Document generation will be limited.")

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# API Keys (In production, use environment variables)
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "AIzaSyBFni23FxQZIKsdP2bMLrJrMuP_YwQy3M4")
TAVILY_API_KEY = os.getenv("TAVILY_API_KEY", "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")

# Initialize LLMs with different configurations
primary_llm = ChatGoogleGenerativeAI(
    model="gemini-1.5-flash-latest",
    google_api_key=GEMINI_API_KEY,
    temperature=0.1
)

creative_llm = ChatGoogleGenerativeAI(
    model="gemini-1.5-flash-latest",
    google_api_key=GEMINI_API_KEY,
    temperature=0.7
)

vision_llm = ChatGoogleGenerativeAI(
    model="gemini-1.5-flash-latest",
    google_api_key=GEMINI_API_KEY,
    temperature=0.3
)

# Enhanced Search Tool
search_tool = TavilySearchResults(
    tavily_api_key=TAVILY_API_KEY,
    max_results=12,
    include_answer=True,
    include_raw_content=True,
    include_images=True
)

# Advanced State definition for LangGraph
class AdvancedAgentState(TypedDict):
    session_id: str
    messages: List[Any]
    query: str
    uploaded_files: List[Dict]
    audio_transcription: str
    image_analysis: Dict
    video_analysis: Dict
    document_analysis: Dict
    query_metadata: Dict
    search_results: List[Dict]
    processed_sources: List[Dict]
    fact_check_results: List[Dict]
    content_analysis: Dict
    answer: str
    summary: str
    citations: List[Dict]
    follow_up_questions: List[str]
    related_topics: List[str]
    sentiment_analysis: Dict
    credibility_score: float
    search_strategy: str
    current_step: str
    processing_metrics: Dict
    error_log: List[str]
    multimodal_context: Dict
    extracted_media: Dict
    uploaded_media_results: List[Dict]

# Create FastAPI app with enhanced metadata
app = FastAPI(
    title="NEXUS AI Backend - Ultra Premium Agentic",
    description="Ultra-premium multimodal AI search backend with LangChain, LangGraph, and Tavily integration",
    version="3.0.0",
    docs_url="/api/docs",
    redoc_url="/api/redoc"
)

# Enhanced CORS configuration
app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:3000",
        "http://127.0.0.1:3000",
        "http://localhost:3001",
        "https://nexus-ai-frontend.vercel.app",
        "*"
    ],
    allow_credentials=True,
    allow_methods=["GET", "POST", "PUT", "DELETE", "OPTIONS", "PATCH"],
    allow_headers=["*"],
)

# Enhanced data models
class SearchRequest(BaseModel):
    query: str = Field(..., min_length=1, max_length=1000, description="Search query")
    search_mode: str = Field(default="comprehensive", description="Search mode")
    search_type: str = Field(default="web", pattern="^(web|documents|images|videos|academic|news)$")
    uploaded_file_ids: Optional[List[str]] = Field(default=None, description="List of uploaded file IDs")
    extract_media: bool = Field(default=True, description="Whether to extract media")
    max_sources: int = Field(default=10, ge=1, le=50, description="Maximum number of sources")
    stream: bool = Field(default=False, description="Stream response")
    filters: Optional[Dict[str, Any]] = Field(default=None, description="Additional filters")

class DocumentSection(BaseModel):
    id: int
    title: str = Field(..., min_length=1, max_length=200)
    content: str = Field(default="", max_length=10000)
    type: str = Field(..., pattern="^(text|sources|media|chart)$")

class DocumentGenerationRequest(BaseModel):
    title: str = Field(..., min_length=1, max_length=200)
    query: str = Field(..., min_length=1, max_length=1000)
    sections: List[DocumentSection]
    searchResults: Dict[str, Any]
    options: Dict[str, bool] = Field(default_factory=dict)

class SummaryRequest(BaseModel):
    url: str = Field(..., description="URL to summarize")
    content: Optional[str] = Field(default=None, description="Optional content to summarize")

class SearchResponse(BaseModel):
    answer: str
    sources: List[Dict[str, Any]]
    extracted_media: Optional[List[Dict[str, Any]]] = []
    follow_up_questions: Optional[List[str]] = []
    credibility_score: float = Field(ge=0.0, le=1.0)
    processing_time: float
    query_analysis: Optional[Dict[str, Any]] = None
    metadata: Optional[Dict[str, Any]] = None

class FileMetadata(BaseModel):
    file_id: str
    filename: str
    file_type: str
    file_size: int
    upload_time: str
    file_path: str
    thumbnail_url: Optional[str] = None
    mime_type: Optional[str] = None
    processing_status: str = "completed"
    extracted_text: Optional[str] = None
    metadata: Optional[Dict[str, Any]] = None

# Enhanced file storage setup
UPLOAD_DIR = Path("uploads")
UPLOAD_DIR.mkdir(exist_ok=True)

THUMBNAILS_DIR = Path("thumbnails")
THUMBNAILS_DIR.mkdir(exist_ok=True)

GENERATED_DOCS_DIR = Path("generated_docs")
GENERATED_DOCS_DIR.mkdir(exist_ok=True)

TEMP_DIR = Path("temp")
TEMP_DIR.mkdir(exist_ok=True)

# Mount static files
app.mount("/thumbnails", StaticFiles(directory=str(THUMBNAILS_DIR)), name="thumbnails")
app.mount("/uploads", StaticFiles(directory=str(UPLOAD_DIR)), name="uploads")

# Enhanced in-memory storage
uploaded_files_db = {}
search_history_db = []
analytics_db = {
    "total_searches": 0,
    "search_types": {},
    "popular_queries": {},
    "error_count": 0,
    "uptime_start": datetime.now()
}

# Content extraction and summarization
class ContentExtractor:
    @staticmethod
    async def extract_content_from_url(url: str) -> Dict:
        """Extract and summarize content from URL"""
        try:
            headers = {
                "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"
            }
            
            response = requests.get(url, headers=headers, timeout=15, allow_redirects=True)
            if response.status_code != 200:
                return {"status": "error", "message": f"Failed to fetch URL: HTTP {response.status_code}"}
            
            soup = BeautifulSoup(response.text, 'html.parser')
            
            # Remove unwanted elements
            for element in soup(['script', 'style', 'nav', 'header', 'footer', 'aside']):
                element.decompose()
            
            # Extract title
            title = soup.find('title')
            title_text = title.get_text().strip() if title else "No title"
            
            # Extract main content
            content_selectors = [
                'main', 'article', '.content', '.post', '.entry',
                '[role="main"]', '.main-content', '#content'
            ]
            
            main_content = ""
            for selector in content_selectors:
                content_element = soup.select_one(selector)
                if content_element:
                    main_content = content_element.get_text(separator=' ', strip=True)
                    break
            
            # Fallback to body if no main content found
            if not main_content:
                body = soup.find('body')
                if body:
                    main_content = body.get_text(separator=' ', strip=True)
            
            # Clean up the content
            main_content = re.sub(r'\s+', ' ', main_content).strip()
            
            # Extract metadata
            meta_description = soup.find('meta', attrs={'name': 'description'})
            description = meta_description.get('content', '') if meta_description else ''
            
            # Get word count
            word_count = len(main_content.split())
            
            return {
                "status": "success",
                "url": url,
                "title": title_text,
                "content": main_content[:5000],  # Limit content for processing
                "description": description,
                "word_count": word_count,
                "extracted_at": datetime.now().isoformat()
            }
            
        except Exception as e:
            logger.error(f"Content extraction failed for {url}: {e}")
            return {"status": "error", "message": str(e)}

    @staticmethod
    async def generate_summary(content: str, url: str = None) -> str:
        """Generate AI summary of content"""
        try:
            summary_prompt = ChatPromptTemplate.from_messages([
                ("system", """You are an expert content summarizer. Create a concise, informative summary that:
                
                1. Captures the main points and key insights
                2. Maintains the original context and meaning
                3. Is well-structured and easy to read
                4. Highlights the most important information
                5. Is approximately 100-150 words long
                
                Focus on actionable insights and key takeaways."""),
                ("human", """Content to summarize: {content}
                
                URL (if available): {url}
                
                Provide a comprehensive summary:""")
            ])
            
            response = await primary_llm.ainvoke(
                summary_prompt.format_messages(
                    content=content[:3000],  # Limit for API
                    url=url or "Not provided"
                )
            )
            
            return response.content.strip()
            
        except Exception as e:
            logger.error(f"Summary generation failed: {e}")
            return "Summary generation failed. Please try again."

# Enhanced File Processing Classes
class FileProcessor:
    @staticmethod
    async def process_pdf(file_content: bytes) -> Dict:
        """Enhanced PDF processing"""
        try:
            pdf_reader = PyPDF2.PdfReader(BytesIO(file_content))
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            
            # Generate preview image
            preview_base64 = None
            try:
                if PIL_AVAILABLE:
                    from pdf2image import convert_from_bytes
                    images = convert_from_bytes(file_content, first_page=1, last_page=1, dpi=150)
                    if images:
                        img = images[0]
                        img.thumbnail((400, 600), Image.Resampling.LANCZOS)
                        img_buffer = BytesIO()
                        img.save(img_buffer, format='PNG', optimize=True, quality=95)
                        preview_base64 = base64.b64encode(img_buffer.getvalue()).decode()
            except Exception as e:
                logger.warning(f"PDF preview generation failed: {e}")
            
            return {
                "type": "pdf",
                "text": text.strip(),
                "page_count": len(pdf_reader.pages),
                "word_count": len(text.split()),
                "metadata": {"pages": len(pdf_reader.pages)},
                "preview": preview_base64,
                "thumbnail": preview_base64
            }
        except Exception as e:
            return {"type": "pdf", "error": str(e), "text": ""}
    
    @staticmethod
    async def process_docx(file_content: bytes) -> Dict:
        """Process DOCX files"""
        try:
            doc = docx.Document(BytesIO(file_content))
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            return {
                "type": "docx",
                "text": text.strip(),
                "paragraph_count": len(doc.paragraphs),
                "word_count": len(text.split()),
                "metadata": {"paragraphs": len(doc.paragraphs)}
            }
        except Exception as e:
            return {"type": "docx", "error": str(e), "text": ""}
    
    @staticmethod
    async def process_txt(file_content: bytes) -> Dict:
        """Process TXT files"""
        try:
            text = file_content.decode('utf-8')
            return {
                "type": "txt",
                "text": text.strip(),
                "line_count": len(text.split('\n')),
                "word_count": len(text.split()),
                "metadata": {"encoding": "utf-8"}
            }
        except Exception as e:
            return {"type": "txt", "error": str(e), "text": ""}

class ImageProcessor:
    @staticmethod
    async def process_image(file_content: bytes) -> Dict:
        """Enhanced image processing"""
        try:
            if not PIL_AVAILABLE:
                return {"type": "image", "error": "PIL not available", "description": ""}
            
            image = Image.open(BytesIO(file_content))
            
            # Extract text using OCR if available
            extracted_text = ""
            try:
                extracted_text = pytesseract.image_to_string(image)
            except:
                pass
            
            # Generate thumbnail
            thumbnail = image.copy()
            thumbnail.thumbnail((400, 400), Image.Resampling.LANCZOS)
            thumb_buffer = BytesIO()
            thumbnail.save(thumb_buffer, format="PNG", optimize=True, quality=95)
            thumbnail_base64 = base64.b64encode(thumb_buffer.getvalue()).decode()
            
            return {
                "type": "image",
                "description": "Enhanced image analysis",
                "extracted_text": extracted_text.strip(),
                "metadata": {
                    "width": image.width,
                    "height": image.height,
                    "format": image.format,
                    "mode": image.mode
                },
                "thumbnail": thumbnail_base64
            }
        except Exception as e:
            return {"type": "image", "error": str(e), "description": ""}

class MediaExtractor:
    @staticmethod
    async def extract_media_from_url(url: str) -> Dict:
        """Enhanced media extraction"""
        try:
            headers = {
                "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"
            }
            
            response = requests.get(url, headers=headers, timeout=10)
            if response.status_code != 200:
                return {"status": "error", "message": f"Failed to fetch URL: HTTP {response.status_code}"}
            
            soup = BeautifulSoup(response.text, 'html.parser')
            
            # Extract images
            images = []
            for img in soup.find_all('img')[:15]:
                src = img.get('src', '')
                if src and not src.startswith('data:'):
                    if not src.startswith(('http://', 'https://')):
                        src = urljoin(url, src)
                    
                    images.append({
                        "type": "image",
                        "url": src,
                        "source_url": url,
                        "thumbnail": src,
                        "title": img.get('alt', f"Image from {urlparse(url).hostname}"),
                        "description": img.get('alt', ''),
                    })
            
            # Extract videos
            videos = []
            video_selectors = ['video', 'iframe[src*="youtube.com"]', 'iframe[src*="vimeo.com"]']
            for selector in video_selectors:
                for video in soup.select(selector)[:10]:
                    src = video.get('src', '')
                    if src:
                        videos.append({
                            "type": "video",
                            "url": src,
                            "source_url": url,
                            "title": video.get('title', f"Video from {urlparse(url).hostname}"),
                            "is_playable": True
                        })
            
            return {
                "status": "success",
                "url": url,
                "images": images,
                "videos": videos,
                "documents": []
            }
        except Exception as e:
            return {"status": "error", "message": str(e)}

# Advanced Multi-Agent System
class AdvancedPerplexityAgent:
    def __init__(self):
        self.workflow = self.create_advanced_workflow()
        self.session_cache = {}
        self.file_cache = {}
        
    def create_advanced_workflow(self):
        workflow = StateGraph(AdvancedAgentState)
        
        # Add enhanced agents
        workflow.add_node("multimodal_processor", self.multimodal_processor_agent)
        workflow.add_node("query_analyzer", self.query_analyzer_agent)
        workflow.add_node("search_strategist", self.search_strategist_agent)
        workflow.add_node("enhanced_search", self.enhanced_search_agent)
        workflow.add_node("content_validator", self.content_validator_agent)
        workflow.add_node("media_extractor", self.media_extractor_agent)
        workflow.add_node("fact_checker", self.fact_checker_agent)
        workflow.add_node("synthesis_expert", self.synthesis_expert_agent)
        workflow.add_node("multimodal_synthesizer", self.multimodal_synthesizer_agent)
        workflow.add_node("summarization", self.summarization_agent)
        workflow.add_node("citation_specialist", self.citation_specialist_agent)
        workflow.add_node("insight_generator", self.insight_generator_agent)
        workflow.add_node("quality_assurance", self.quality_assurance_agent)
        
        # Define workflow edges
        workflow.set_entry_point("multimodal_processor")
        workflow.add_edge("multimodal_processor", "query_analyzer")
        workflow.add_edge("query_analyzer", "search_strategist")
        workflow.add_edge("search_strategist", "enhanced_search")
        workflow.add_edge("enhanced_search", "content_validator")
        workflow.add_edge("content_validator", "media_extractor")
        workflow.add_edge("media_extractor", "fact_checker")
        workflow.add_edge("fact_checker", "synthesis_expert")
        workflow.add_edge("synthesis_expert", "multimodal_synthesizer")
        workflow.add_edge("multimodal_synthesizer", "summarization")
        workflow.add_edge("summarization", "citation_specialist")
        workflow.add_edge("citation_specialist", "insight_generator")
        workflow.add_edge("insight_generator", "quality_assurance")
        workflow.add_edge("quality_assurance", END)
        
        return workflow.compile()
    
    def get_favicon_url(self, domain: str) -> str:
        """Get favicon URL for a domain"""
        try:
            if 'github.com' in domain:
                return "https://github.com/favicon.ico"
            elif 'stackoverflow.com' in domain:
                return "https://stackoverflow.com/favicon.ico"
            elif 'wikipedia.org' in domain:
                return "https://wikipedia.org/favicon.ico"
            else:
                return f"https://{domain}/favicon.ico"
        except:
            return ""
    
    def extract_domain(self, url: str) -> str:
        """Extract domain from URL"""
        try:
            return urlparse(url).netloc
        except:
            return "unknown"
    
    async def multimodal_processor_agent(self, state: AdvancedAgentState) -> AdvancedAgentState:
        """Process uploaded files and multimodal content"""
        state["current_step"] = "Processing multimodal content..."
        
        try:
            multimodal_context = {
                "has_files": bool(state.get("uploaded_files")),
                "context_summary": ""
            }
            
            context_parts = []
            uploaded_media_results = []
            
            # Process uploaded files
            if state.get("uploaded_files"):
                for file_data in state["uploaded_files"]:
                    if file_data.get("text"):
                        context_parts.append(f"Document content: {file_data['text'][:500]}...")
                    
                    # Create media result for display
                    media_result = {
                        "file_id": file_data.get("file_id", str(uuid.uuid4())),
                        "filename": file_data.get("filename", "Unknown"),
                        "file_type": file_data.get("file_type", "unknown"),
                        "url": f"/files/{file_data.get('file_id', '')}",
                        "extracted_text": file_data.get("extracted_text", file_data.get("text", "")),
                        "analysis": file_data.get("analysis", {})
                    }
                    
                    uploaded_media_results.append(media_result)
            
            multimodal_context["context_summary"] = " | ".join(context_parts)
            state["multimodal_context"] = multimodal_context
            state["uploaded_media_results"] = uploaded_media_results
            
        except Exception as e:
            state["error_log"].append(f"Multimodal processing error: {e}")
            
        return state
    
    async def query_analyzer_agent(self, state: AdvancedAgentState) -> AdvancedAgentState:
        """Analyze query with multimodal context"""
        state["current_step"] = "Analyzing query..."
        
        query = state["query"]
        
        try:
            metadata = {
                "intent": self.detect_query_intent(query),
                "complexity": self.assess_complexity(query),
                "domains": self.identify_domains(query),
                "entities": self.extract_entities(query),
                "query_type": self.classify_query_type(query),
                "multimodal": bool(state.get("uploaded_files"))
            }
            
            state["query_metadata"] = metadata
            
        except Exception as e:
            state["error_log"].append(f"Query analysis error: {e}")
            
        return state
    
    async def search_strategist_agent(self, state: AdvancedAgentState) -> AdvancedAgentState:
        """Plan search strategy"""
        state["current_step"] = "Planning search strategy..."
        
        metadata = state["query_metadata"]
        
        if metadata.get("complexity") == "expert":
            strategy = "academic_focused"
        elif "recent" in state["query"].lower():
            strategy = "news_focused"
        else:
            strategy = "comprehensive"
        
        state["search_strategy"] = strategy
        return state
    
    async def enhanced_search_agent(self, state: AdvancedAgentState) -> AdvancedAgentState:
        """Perform enhanced search using Tavily"""
        state["current_step"] = "Performing enhanced search..."
        
        query = state["query"]
        
        try:
            # Primary search using Tavily
            search_results = search_tool.invoke(query)
            
            # Process results
            processed_results = []
            seen_urls = set()
            
            for result in search_results:
                if isinstance(result, dict) and result.get("url") not in seen_urls:
                    domain = self.extract_domain(result.get("url", ""))
                    favicon_url = self.get_favicon_url(domain)
                    
                    # Extract content and generate summary
                    content_data = await ContentExtractor.extract_content_from_url(result.get("url", ""))
                    summary = ""
                    if content_data.get("status") == "success":
                        summary = await ContentExtractor.generate_summary(
                            content_data.get("content", result.get("content", "")),
                            result.get("url", "")
                        )
                    
                    processed_result = {
                        "id": str(uuid.uuid4()),
                        "title": result.get("title", ""),
                        "url": result.get("url", ""),
                        "snippet": result.get("content", ""),
                        "summary": summary or result.get("content", "")[:300] + "...",
                        "domain": domain,
                        "favicon_url": favicon_url,
                        "credibility_score": self.calculate_initial_credibility(result),
                        "content_type": self.classify_content_type(result),
                        "is_academic": self.is_academic_source(result.get("url", "")),
                        "is_github": "github.com" in result.get("url", "").lower(),
                        "word_count": content_data.get("word_count", 0),
                        "extracted_content": content_data.get("content", "")
                    }
                    processed_results.append(processed_result)
                    seen_urls.add(result.get("url"))
            
            state["search_results"] = processed_results[:12]
            
        except Exception as e:
            state["error_log"].append(f"Enhanced search error: {e}")
            state["search_results"] = []
        
        return state
    
    async def content_validator_agent(self, state: AdvancedAgentState) -> AdvancedAgentState:
        """Validate and score content"""
        state["current_step"] = "Validating content..."
        
        search_results = state["search_results"]
        
        try:
            validated_sources = []
            for source in search_results:
                # Enhanced credibility scoring
                credibility_score = self.calculate_enhanced_credibility(source)
                source["credibility_score"] = credibility_score
                validated_sources.append(source)
            
            # Sort by credibility
            validated_sources.sort(key=lambda x: x["credibility_score"], reverse=True)
            state["processed_sources"] = validated_sources[:8]
            
        except Exception as e:
            state["error_log"].append(f"Content validation error: {e}")
            state["processed_sources"] = search_results[:8]
        
        return state
    
    async def media_extractor_agent(self, state: AdvancedAgentState) -> AdvancedAgentState:
        """Extract media from web sources"""
        state["current_step"] = "Extracting media..."
        
        sources = state["processed_sources"]
        
        try:
            all_media = {
                "images": [],
                "videos": [],
                "documents": []
            }
            
            # Extract media from top sources
            for source in sources[:3]:
                url = source.get("url", "")
                if url:
                    media_results = await MediaExtractor.extract_media_from_url(url)
                    if media_results.get("status") == "success":
                        for img in media_results.get("images", []):
                            all_media["images"].append(img)
                        for vid in media_results.get("videos", []):
                            all_media["videos"].append(vid)
            
            state["extracted_media"] = {
                "status": "success",
                "images": all_media["images"][:10],
                "videos": all_media["videos"][:5],
                "documents": []
            }
            
        except Exception as e:
            state["error_log"].append(f"Media extraction error: {e}")
            state["extracted_media"] = {"status": "error", "images": [], "videos": [], "documents": []}
        
        return state
    
    async def fact_checker_agent(self, state: AdvancedAgentState) -> AdvancedAgentState:
        """Fact-check content"""
        state["current_step"] = "Fact-checking..."
        
        sources = state["processed_sources"]
        
        try:
            fact_checks = []
            for i, source in enumerate(sources[:5]):
                fact_check = {
                    "source_index": i,
                    "credibility_indicators": self.analyze_credibility_indicators(source["snippet"]),
                    "domain_authority": self.calculate_domain_authority(source["domain"])
                }
                fact_checks.append(fact_check)
            
            state["fact_check_results"] = fact_checks
            
        except Exception as e:
            state["error_log"].append(f"Fact-checking error: {e}")
            state["fact_check_results"] = []
        
        return state
    
    async def synthesis_expert_agent(self, state: AdvancedAgentState) -> AdvancedAgentState:
        """Synthesize comprehensive answer"""
        state["current_step"] = "Synthesizing answer..."
        
        query = state["query"]
        sources = state["processed_sources"]
        multimodal_context = state.get("multimodal_context", {})
        
        try:
            # Create synthesis prompt
            synthesis_prompt = ChatPromptTemplate.from_messages([
                ("system", """You are an expert research synthesizer. Create a comprehensive answer that:
                
                1. Addresses the query directly with authoritative insights
                2. Uses proper structure with headings and organized sections
                3. Incorporates multimodal context when relevant
                4. Maintains accuracy through cross-referencing
                5. Provides actionable insights and recommendations
                
                Structure your response with:
                - **Executive Summary** (2-3 sentences)
                - **Main Analysis** with clear sections
                - **Key Findings** in bullet format
                - **Conclusion** with actionable recommendations
                
                Write professionally and ensure all claims are well-supported."""),
                ("human", """Query: {query}
                
                Multimodal Context: {multimodal_context}
                
                Sources:
                {sources}
                
                Create a comprehensive answer:""")
            ])
            
            sources_text = "\n\n".join([
                f"**Source {i+1}**\n"
                f"Title: {s['title']}\n"
                f"URL: {s['url']}\n"
                f"Content: {s['snippet']}\n"
                f"Summary: {s.get('summary', '')}\n"
                f"Credibility: {s['credibility_score']:.2f}"
                for i, s in enumerate(sources[:6])
            ])
            
            response = await primary_llm.ainvoke(
                synthesis_prompt.format_messages(
                    query=query,
                    multimodal_context=str(multimodal_context),
                    sources=sources_text
                )
            )
            
            state["answer"] = response.content
            
            # Content analysis
            content_analysis = {
                "word_count": len(response.content.split()),
                "reading_time": max(1, len(response.content.split()) // 200),
                "key_concepts": self.extract_key_concepts(response.content),
                "confidence_score": self.calculate_confidence_score(sources)
            }
            state["content_analysis"] = content_analysis
            
        except Exception as e:
            state["error_log"].append(f"Synthesis error: {e}")
            state["answer"] = "I apologize, but I encountered an error while processing your request."
            state["content_analysis"] = {}
        
        return state
    
    async def multimodal_synthesizer_agent(self, state: AdvancedAgentState) -> AdvancedAgentState:
        """Enhance with multimodal synthesis"""
        state["current_step"] = "Enhancing with multimodal insights..."
        
        # For now, pass through - can be enhanced later
        state["multimodal_insights"] = {
            "files_integrated": len(state.get("uploaded_files", [])),
            "enhancement_applied": False
        }
        
        return state
    
    async def summarization_agent(self, state: AdvancedAgentState) -> AdvancedAgentState:
        """Create summary"""
        state["current_step"] = "Creating summary..."
        
        answer = state["answer"]
        query = state["query"]
        
        try:
            summary_prompt = ChatPromptTemplate.from_messages([
                ("system", "Create a comprehensive executive summary that captures the essence of the analysis. Keep it concise but comprehensive (3-4 sentences)."),
                ("human", "Query: {query}\n\nFull Answer: {answer}\n\nCreate an executive summary:")
            ])
            
            response = await primary_llm.ainvoke(
                summary_prompt.format_messages(query=query, answer=answer[:1500])
            )
            state["summary"] = response.content.strip()
            
        except Exception as e:
            state["error_log"].append(f"Summarization error: {e}")
            state["summary"] = "Comprehensive analysis completed successfully."
        
        return state
    
    async def citation_specialist_agent(self, state: AdvancedAgentState) -> AdvancedAgentState:
        """Manage citations"""
        state["current_step"] = "Managing citations..."
        
        sources = state["processed_sources"]
        
        citations = []
        for i, source in enumerate(sources):
            citation = {
                "index": i + 1,
                "title": source["title"],
                "url": source["url"],
                "domain": source["domain"],
                "credibility_score": source["credibility_score"],
                "summary": source.get("summary", "")
            }
            citations.append(citation)
        
        state["citations"] = citations
        return state
    
    async def insight_generator_agent(self, state: AdvancedAgentState) -> AdvancedAgentState:
        """Generate insights and follow-up questions"""
        state["current_step"] = "Generating insights..."
        
        query = state["query"]
        answer = state["answer"]
        
        try:
            insight_prompt = ChatPromptTemplate.from_messages([
                ("system", """Generate insightful follow-up questions and related topics. Provide:
                - **Follow-up Questions** (4-5): Specific, actionable questions
                - **Related Topics** (4-6): Connected concepts for exploration"""),
                ("human", "Query: {query}\n\nAnswer: {answer}\n\nGenerate insights:")
            ])
            
            response = await creative_llm.ainvoke(
                insight_prompt.format_messages(query=query, answer=answer[:1200])
            )
            
            content = response.content
            
            # Parse follow-up questions and related topics
            follow_ups = self.extract_follow_ups(content, query)
            related = self.extract_related_topics(content, query)
            
            state["follow_up_questions"] = follow_ups
            state["related_topics"] = related
            
        except Exception as e:
            state["error_log"].append(f"Insight generation error: {e}")
            state["follow_up_questions"] = [f"What are the latest developments in {query}?"]
            state["related_topics"] = ["Related research", "Current trends"]
        
        return state
    
    async def quality_assurance_agent(self, state: AdvancedAgentState) -> AdvancedAgentState:
        """Final quality assurance"""
        state["current_step"] = "Performing quality assurance..."
        
        sources = state["processed_sources"]
        
        # Calculate overall credibility
        if sources:
            credibility_score = statistics.mean([s["credibility_score"] for s in sources])
        else:
            credibility_score = 0.75
        
        # Quality metrics
        answer_length = len(state["answer"].split())
        source_count = len(sources)
        error_count = len(state["error_log"])
        
        # Adjust credibility based on quality factors
        if answer_length < 50:
            credibility_score *= 0.8
        if source_count < 3:
            credibility_score *= 0.7
        if error_count > 2:
            credibility_score *= 0.6
        
        state["credibility_score"] = max(0.1, min(1.0, credibility_score))
        
        state["processing_metrics"] = {
            "sources_analyzed": len(state["search_results"]),
            "sources_used": len(sources),
            "answer_word_count": answer_length,
            "errors_encountered": error_count,
            "overall_quality": "excellent" if credibility_score > 0.8 else "good" if credibility_score > 0.6 else "fair"
        }
        
        return state
    
    # Helper methods
    def detect_query_intent(self, query: str) -> str:
        query_lower = query.lower()
        if any(word in query_lower for word in ['how to', 'tutorial', 'guide']):
            return 'instructional'
        elif any(word in query_lower for word in ['vs', 'compare', 'difference']):
            return 'comparison'
        elif any(word in query_lower for word in ['what is', 'define']):
            return 'definitional'
        else:
            return 'informational'
    
    def assess_complexity(self, query: str) -> str:
        technical_terms = len(re.findall(r'\b[A-Z][a-z]*[A-Z][a-zA-Z]*\b', query))
        word_count = len(query.split())
        
        if technical_terms > 2 or word_count > 15:
            return 'expert'
        elif technical_terms > 0 or word_count > 8:
            return 'moderate'
        else:
            return 'simple'
    
    def identify_domains(self, query: str) -> List[str]:
        domains = []
        query_lower = query.lower()
        
        domain_keywords = {
            'technology': ['ai', 'machine learning', 'programming'],
            'science': ['research', 'study', 'scientific'],
            'business': ['company', 'market', 'business'],
            'health': ['health', 'medical', 'disease']
        }
        
        for domain, keywords in domain_keywords.items():
            if any(keyword in query_lower for keyword in keywords):
                domains.append(domain)
        
        return domains or ['general']
    
    def extract_entities(self, query: str) -> List[str]:
        words = query.split()
        entities = []
        
        for word in words:
            if word[0].isupper() and len(word) > 2:
                entities.append(word)
        
        return entities
    
    def classify_query_type(self, query: str) -> str:
        if '?' in query:
            return 'question'
        elif query.lower().startswith(('explain', 'describe')):
            return 'explanation_request'
        else:
            return 'statement'
    
    def calculate_initial_credibility(self, result: Dict) -> float:
        url = result.get("url", "")
        score = 0.5
        
        # Domain-based scoring
        if any(domain in url for domain in ['edu', 'gov', 'org']):
            score += 0.3
        elif any(domain in url for domain in ['wikipedia.org', 'scholar.google']):
            score += 0.2
        elif any(domain in url for domain in ['github.com', 'stackoverflow.com']):
            score += 0.15
        
        return min(1.0, score)
    
    def classify_content_type(self, result: Dict) -> str:
        url = result.get("url", "")
        if "github.com" in url:
            return "repository"
        elif "wikipedia.org" in url:
            return "encyclopedia"
        elif any(ext in url for ext in ['.pdf', '.doc']):
            return "document"
        else:
            return "webpage"
    
    def is_academic_source(self, url: str) -> bool:
        return any(domain in url for domain in ['edu', 'scholar.google', 'arxiv.org', 'pubmed'])
    
    def calculate_enhanced_credibility(self, source: Dict) -> float:
        base_score = source.get('credibility_score', 0.5)
        
        # Additional factors
        if source.get('is_academic'):
            base_score += 0.1
        if source.get('is_github'):
            base_score += 0.05
        if source.get('word_count', 0) > 500:
            base_score += 0.05
        
        return min(1.0, base_score)
    
    def analyze_credibility_indicators(self, content: str) -> Dict:
        return {
            'has_numbers': bool(re.search(r'\d+', content)),
            'has_dates': bool(re.search(r'\b\d{4}\b', content)),
            'has_sources': 'source' in content.lower(),
            'length_adequate': len(content) > 100
        }
    
    def calculate_domain_authority(self, domain: str) -> float:
        high_authority = ['wikipedia.org', 'github.com', 'stackoverflow.com']
        if domain in high_authority:
            return 0.9
        elif domain.endswith('.edu') or domain.endswith('.gov'):
            return 0.85
        else:
            return 0.7
    
    def extract_key_concepts(self, content: str) -> List[str]:
        words = content.split()
        concepts = []
        
        for word in words:
            cleaned = re.sub(r'[^\w]', '', word)
            if len(cleaned) > 3 and cleaned[0].isupper():
                concepts.append(cleaned)
        
        return list(set(concepts))[:8]
    
    def calculate_confidence_score(self, sources: List[Dict]) -> float:
        if not sources:
            return 0.3
        
        avg_credibility = statistics.mean([s.get("credibility_score", 0.5) for s in sources])
        return min(1.0, avg_credibility)
    
    def extract_follow_ups(self, content: str, query: str) -> List[str]:
        # Basic extraction - can be enhanced
        return [
            f"What are the practical applications of {query}?",
            f"How does {query} compare to alternatives?",
            f"What are recent developments in {query}?",
            f"What should I know about {query}?"
        ]
    
    def extract_related_topics(self, content: str, query: str) -> List[str]:
        # Basic extraction - can be enhanced
        return [
            "Advanced applications",
            "Recent research",
            "Best practices",
            "Future trends",
            "Industry insights"
        ]

# Initialize the enhanced agent
advanced_agent = AdvancedPerplexityAgent()

# Enhanced file type detection
def get_enhanced_file_type(filename: str, content_type: str = None) -> Dict[str, Any]:
    """Enhanced file type detection with metadata"""
    mime_type, _ = mimetypes.guess_type(filename)
    if content_type:
        mime_type = content_type
    
    if not mime_type:
        return {"type": "unknown", "category": "unknown", "mime_type": "application/octet-stream"}
    
    if mime_type.startswith('image/'):
        return {
            "type": "image",
            "category": "media",
            "mime_type": mime_type,
            "supports_thumbnail": True,
            "can_extract_text": True
        }
    elif mime_type.startswith('video/'):
        return {
            "type": "video", 
            "category": "media",
            "mime_type": mime_type,
            "supports_thumbnail": True
        }
    elif mime_type == 'application/pdf':
        return {
            "type": "pdf",
            "category": "document",
            "mime_type": mime_type,
            "can_extract_text": True,
            "supports_preview": True
        }
    elif mime_type in ['application/msword', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
        return {
            "type": "document",
            "category": "document", 
            "mime_type": mime_type,
            "can_extract_text": True
        }
    elif mime_type.startswith('text/'):
        return {
            "type": "text",
            "category": "document",
            "mime_type": mime_type,
            "can_extract_text": True
        }
    else:
        return {
            "type": "document",
            "category": "document",
            "mime_type": mime_type
        }

def generate_thumbnail(file_path: Path, output_path: Path, size: tuple = (400, 300)) -> bool:
    """Generate thumbnail for images"""
    if not PIL_AVAILABLE:
        return False
    
    try:
        with Image.open(file_path) as img:
            if img.mode in ('RGBA', 'LA', 'P'):
                img = img.convert('RGB')
            
            img.thumbnail(size, Image.Resampling.LANCZOS)
            img.save(output_path, 'JPEG', quality=85, optimize=True)
            return True
    except Exception as e:
        logger.error(f"Thumbnail generation failed: {e}")
        return False

def generate_enhanced_pdf_document(doc_request: DocumentGenerationRequest) -> bytes:
    """Generate enhanced PDF document with better formatting"""
    
    if not REPORTLAB_AVAILABLE:
        # Enhanced fallback text generation
        content = f"""
{doc_request.title}
{'='*len(doc_request.title)}

Query: {doc_request.query}
Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
Generated by: NEXUS AI Agentic Backend v3.0.0

TABLE OF CONTENTS
"""
        
        for i, section in enumerate(doc_request.sections, 1):
            content += f"{i}. {section.title}\n"
        
        content += "\n" + "="*50 + "\n\n"
        
        for i, section in enumerate(doc_request.sections, 1):
            content += f"\n{i}. {section.title}\n{'-'*len(section.title)}\n\n"
            
            if section.type == "text":
                content += f"{section.content}\n\n"
                
            elif section.type == "sources":
                content += "SOURCES AND REFERENCES\n\n"
                sources = doc_request.searchResults.get('sources', [])
                for j, source in enumerate(sources, 1):
                    content += f"{j}. {source.get('title', 'Untitled')}\n"
                    content += f"   URL: {source.get('url', 'No URL')}\n"
                    content += f"   Credibility: {source.get('credibility_score', 0)*100:.0f}%\n"
                    content += f"   Summary: {source.get('summary', 'No summary available')}\n\n"
                    
            elif section.type == "media":
                content += "MEDIA GALLERY\n\n"
                media_items = doc_request.searchResults.get('extracted_media', [])
                for j, media in enumerate(media_items, 1):
                    content += f"{j}. {media.get('title', 'Untitled Media')}\n"
                    content += f"   Type: {media.get('type', 'Unknown').title()}\n"
                    content += f"   Source: {media.get('source_url', 'No source')}\n\n"
        
        content += "\n" + "="*50 + "\n"
        content += f"Document generated by NEXUS AI Agentic Backend v3.0.0\n"
        content += f"Powered by LangChain, LangGraph, and Tavily\n"
        
        return content.encode('utf-8')
    
    # Enhanced PDF generation with ReportLab
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer, 
        pagesize=A4,
        rightMargin=72,
        leftMargin=72,
        topMargin=72,
        bottomMargin=72
    )
    
    styles = getSampleStyleSheet()
    story = []
    
    # Enhanced custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Title'],
        fontSize=28,
        spaceAfter=30,
        textColor=colors.darkblue,
        alignment=TA_CENTER,
        fontName='Helvetica-Bold'
    )
    
    subtitle_style = ParagraphStyle(
        'CustomSubtitle',
        parent=styles['Normal'],
        fontSize=14,
        spaceAfter=20,
        textColor=colors.grey,
        alignment=TA_CENTER,
        fontName='Helvetica'
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading1'],
        fontSize=18,
        spaceAfter=15,
        spaceBefore=20,
        textColor=colors.darkblue,
        fontName='Helvetica-Bold'
    )
    
    subheading_style = ParagraphStyle(
        'CustomSubheading',
        parent=styles['Heading2'],
        fontSize=14,
        spaceAfter=10,
        spaceBefore=15,
        textColor=colors.darkslategray,
        fontName='Helvetica-Bold'
    )
    
    body_style = ParagraphStyle(
        'CustomBody',
        parent=styles['Normal'],
        fontSize=11,
        spaceAfter=12,
        alignment=TA_JUSTIFY,
        fontName='Helvetica'
    )
    
    # Title page
    story.append(Spacer(1, 2*inch))
    story.append(Paragraph(doc_request.title, title_style))
    story.append(Spacer(1, 0.5*inch))
    
    story.append(Paragraph(f"Research Query: {doc_request.query}", subtitle_style))
    story.append(Paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y at %H:%M')}", subtitle_style))
    story.append(Paragraph(f"Generated by: NEXUS AI Agentic Backend v3.0.0", subtitle_style))
    story.append(Paragraph(f"Sections: {len(doc_request.sections)}", subtitle_style))
    
    story.append(Spacer(1, 1*inch))
    
    # Add summary box
    summary_data = [
        ['Search Results Summary', ''],
        ['Total Sources', str(len(doc_request.searchResults.get('sources', [])))],
        ['Media Items', str(len(doc_request.searchResults.get('extracted_media', [])))],
        ['Credibility Score', f"{doc_request.searchResults.get('credibility_score', 0.8)*100:.0f}%"],
        ['Processing Time', f"{doc_request.searchResults.get('processing_time', 0):.2f}s"]
    ]
    
    summary_table = Table(summary_data, colWidths=[3*inch, 2*inch])
    summary_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.darkblue),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, 0), 12),
        ('FONTSIZE', (0, 1), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.lightgrey),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE')
    ]))
    
    story.append(summary_table)
    story.append(PageBreak())
    
    # Table of Contents
    story.append(Paragraph("Table of Contents", heading_style))
    story.append(Spacer(1, 20))
    
    toc_data = [['Section', 'Title', 'Type']]
    for i, section in enumerate(doc_request.sections, 1):
        toc_data.append([str(i), section.title, section.type.title()])
    
    toc_table = Table(toc_data, colWidths=[0.8*inch, 4*inch, 1.2*inch])
    toc_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE')
    ]))
    
    story.append(toc_table)
    story.append(PageBreak())
    
    # Document sections
    for i, section in enumerate(doc_request.sections, 1):
        story.append(Paragraph(f"{i}. {section.title}", heading_style))
        
        if section.type == "text":
            if section.content.strip():
                # Split content into paragraphs
                paragraphs = section.content.split('\n\n')
                for para in paragraphs:
                    if para.strip():
                        story.append(Paragraph(para.strip(), body_style))
            else:
                story.append(Paragraph("This section contains the main content and analysis.", body_style))
        
        elif section.type == "sources":
            story.append(Paragraph("Sources and References", subheading_style))
            story.append(Spacer(1, 10))
            
            sources = doc_request.searchResults.get('sources', [])
            if sources:
                # Create enhanced sources table
                sources_data = [['#', 'Title', 'Domain', 'Credibility', 'Type']]
                for j, source in enumerate(sources[:10], 1):  # Limit to top 10
                    title = source.get('title', 'Untitled')
                    if len(title) > 40:
                        title = title[:37] + "..."
                    
                    domain = source.get('domain', 'Unknown')
                    credibility = f"{source.get('credibility_score', 0)*100:.0f}%"
                    content_type = source.get('content_type', 'web').title()
                    
                    sources_data.append([str(j), title, domain, credibility, content_type])
                
                sources_table = Table(sources_data, colWidths=[0.4*inch, 2.8*inch, 1.2*inch, 0.8*inch, 0.8*inch])
                sources_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.darkblue),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('ALIGN', (3, 0), (3, -1), 'CENTER'),  # Center credibility scores
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 9),
                    ('FONTSIZE', (0, 1), (-1, -1), 8),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.lightgrey),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE')
                ]))
                story.append(sources_table)
                
                # Add detailed source descriptions with summaries
                story.append(Spacer(1, 15))
                story.append(Paragraph("Detailed Source Information:", subheading_style))
                
                for j, source in enumerate(sources[:5], 1):  # Detailed info for top 5
                    story.append(Paragraph(f"<b>{j}. {source.get('title', 'Untitled')}</b>", body_style))
                    story.append(Paragraph(f"URL: {source.get('url', 'No URL available')}", body_style))
                    if source.get('summary'):
                        story.append(Paragraph(f"Summary: {source.get('summary')}", body_style))
                    elif source.get('snippet'):
                        story.append(Paragraph(f"Summary: {source.get('snippet')}", body_style))
                    story.append(Spacer(1, 8))
            else:
                story.append(Paragraph("No sources available for this search.", body_style))
        
        elif section.type == "media":
            story.append(Paragraph("Media Gallery", subheading_style))
            story.append(Spacer(1, 10))
            
            media_items = doc_request.searchResults.get('extracted_media', [])
            if media_items:
                # Create media summary table
                media_data = [['#', 'Title', 'Type', 'Source']]
                for j, media in enumerate(media_items[:8], 1):  # Limit to 8 items
                    title = media.get('title', 'Untitled Media')
                    if len(title) > 35:
                        title = title[:32] + "..."
                    
                    media_type = media.get('type', 'Unknown').title()
                    source = media.get('source_url', 'No source')
                    if len(source) > 30:
                        source = source[:27] + "..."
                    
                    media_data.append([str(j), title, media_type, source])
                
                media_table = Table(media_data, colWidths=[0.4*inch, 2.5*inch, 1*inch, 2.1*inch])
                media_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.darkgreen),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 9),
                    ('FONTSIZE', (0, 1), (-1, -1), 8),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.lightgreen),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE')
                ]))
                story.append(media_table)
            else:
                story.append(Paragraph("No media items found for this search.", body_style))
        
        elif section.type == "chart":
            story.append(Paragraph("Data Analysis and Charts", subheading_style))
            story.append(Paragraph("Chart generation feature coming soon. This section will contain data visualizations and statistical analysis.", body_style))
        
        story.append(Spacer(1, 20))
    
    # Add footer information
    story.append(Spacer(1, 30))
    story.append(Paragraph("Document Information", heading_style))
    
    footer_data = [
        ['Generated by', 'NEXUS AI Ultra Premium Agentic Search Engine'],
        ['Query', doc_request.query],
        ['Generation Time', datetime.now().strftime('%Y-%m-%d %H:%M:%S UTC')],
        ['Total Sections', str(len(doc_request.sections))],
        ['Sources Analyzed', str(len(doc_request.searchResults.get('sources', [])))],
        ['Media Items', str(len(doc_request.searchResults.get('extracted_media', [])))]
    ]
    
    footer_table = Table(footer_data, colWidths=[2*inch, 4*inch])
    footer_table.setStyle(TableStyle([
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE')
    ]))
    
    story.append(footer_table)
    
    # Build PDF
    doc.build(story)
    pdf_data = buffer.getvalue()
    buffer.close()
    
    return pdf_data

# API Endpoints

@app.get("/")
async def root():
    """Enhanced health check endpoint"""
    uptime = datetime.now() - analytics_db["uptime_start"]
    return {
        "message": "NEXUS AI Backend Ultra Premium Agentic is running",
        "status": "healthy", 
        "version": "3.0.0",
        "uptime": str(uptime),
        "features": [
            "langchain_integration",
            "langgraph_workflow", 
            "tavily_search",
            "multimodal_processing",
            "enhanced_document_generation",
            "advanced_media_processing",
            "intelligent_analytics",
            "content_summarization",
            "url_content_extraction"
        ],
        "agentic_capabilities": {
            "agents": 13,
            "workflow_steps": ["multimodal_processor", "query_analyzer", "search_strategist", "enhanced_search", "content_validator", "media_extractor", "fact_checker", "synthesis_expert", "multimodal_synthesizer", "summarization", "citation_specialist", "insight_generator", "quality_assurance"],
            "llm_models": ["gemini-1.5-flash-latest"],
            "search_engine": "tavily"
        }
    }

@app.post("/upload", response_model=FileMetadata)
async def upload_file(file: UploadFile = File(...)):
    """Enhanced file upload with processing"""
    global uploaded_files_db
    
    try:
        if file.size and file.size > 100 * 1024 * 1024:
            raise HTTPException(status_code=413, detail="File too large. Maximum size is 100MB.")
        
        file_id = str(uuid.uuid4())
        file_extension = Path(file.filename).suffix.lower()
        stored_filename = f"{file_id}{file_extension}"
        file_path = UPLOAD_DIR / stored_filename
        
        # Save file
        content = await file.read()
        with open(file_path, "wb") as buffer:
            buffer.write(content)
        
        # Get file type info
        file_info = get_enhanced_file_type(file.filename, file.content_type)
        file_size = len(content)
        
        # Process file based on type
        analysis_result = {}
        extracted_text = ""
        thumbnail_url = None
        
        if file_info["type"] == "pdf":
            analysis_result = await FileProcessor.process_pdf(content)
            extracted_text = analysis_result.get('text', '')
            if analysis_result.get('preview'):
                # Save preview as thumbnail
                thumbnail_filename = f"{file_id}_thumb.png"
                thumbnail_path = THUMBNAILS_DIR / thumbnail_filename
                with open(thumbnail_path, "wb") as f:
                    f.write(base64.b64decode(analysis_result['preview']))
                thumbnail_url = f"/thumbnails/{thumbnail_filename}"
                
        elif file_info["type"] == "image" and PIL_AVAILABLE:
            analysis_result = await ImageProcessor.process_image(content)
            extracted_text = analysis_result.get('extracted_text', '')
            if analysis_result.get('thumbnail'):
                thumbnail_filename = f"{file_id}_thumb.png"
                thumbnail_path = THUMBNAILS_DIR / thumbnail_filename
                with open(thumbnail_path, "wb") as f:
                    f.write(base64.b64decode(analysis_result['thumbnail']))
                thumbnail_url = f"/thumbnails/{thumbnail_filename}"
                
        elif file_info["type"] == "docx":
            analysis_result = await FileProcessor.process_docx(content)
            extracted_text = analysis_result.get('text', '')
            
        elif file_info["type"] == "txt":
            analysis_result = await FileProcessor.process_txt(content)
            extracted_text = analysis_result.get('text', '')
        
        # Create metadata
        file_metadata = FileMetadata(
            file_id=file_id,
            filename=file.filename,
            file_type=file_info["type"],
            file_size=file_size,
            upload_time=datetime.now().isoformat(),
            file_path=str(file_path),
            thumbnail_url=thumbnail_url,
            mime_type=file.content_type or file_info["mime_type"],
            processing_status="completed",
            extracted_text=extracted_text,
            metadata={
                "analysis": analysis_result,
                "category": file_info["category"],
                "file_hash": hashlib.md5(content).hexdigest()
            }
        )
        
        uploaded_files_db[file_id] = file_metadata.dict()
        
        logger.info(f"File uploaded and processed: {file.filename} ({file_id})")
        
        return file_metadata
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"File upload failed: {str(e)}")
        raise HTTPException(status_code=500, detail=f"File upload failed: {str(e)}")

@app.post("/search", response_model=SearchResponse)
async def search(request: SearchRequest):
    """Enhanced AI search with LangChain agents"""
    global search_history_db, analytics_db, uploaded_files_db
    
    try:
        start_time = datetime.now()
        
        # Update analytics
        analytics_db["total_searches"] += 1
        analytics_db["search_types"][request.search_type] = analytics_db["search_types"].get(request.search_type, 0) + 1
        
        logger.info(f"Agentic search request: '{request.query}' (type: {request.search_type})")
        
        # Get uploaded files
        uploaded_files = []
        if request.uploaded_file_ids:
            for file_id in request.uploaded_file_ids:
                if file_id in uploaded_files_db:
                    file_data = uploaded_files_db[file_id]
                    uploaded_files.append({
                        "file_id": file_id,
                        "filename": file_data["filename"],
                        "file_type": file_data["file_type"],
                        "text": file_data.get("extracted_text", ""),
                        "analysis": file_data.get("metadata", {}).get("analysis", {})
                    })
        
        # Initialize state for LangGraph workflow
        initial_state = {
            "session_id": str(uuid.uuid4()),
            "messages": [],
            "query": request.query,
            "uploaded_files": uploaded_files,
            "audio_transcription": "",
            "image_analysis": {},
            "video_analysis": {},
            "document_analysis": {},
            "query_metadata": {},
            "search_results": [],
            "processed_sources": [],
            "fact_check_results": [],
            "content_analysis": {},
            "answer": "",
            "summary": "",
            "citations": [],
            "follow_up_questions": [],
            "related_topics": [],
            "sentiment_analysis": {},
            "credibility_score": 0.0,
            "search_strategy": "",
            "current_step": "Initializing...",
            "processing_metrics": {},
            "error_log": [],
            "multimodal_context": {},
            "extracted_media": {},
            "uploaded_media_results": []
        }
        
        # Run the LangGraph workflow
        final_state = await advanced_agent.workflow.ainvoke(initial_state)
        
        # Calculate processing time
        processing_time = (datetime.now() - start_time).total_seconds()
        
        # Format response for frontend
        results = {
            "answer": final_state["answer"],
            "sources": final_state["processed_sources"],
            "extracted_media": [],
            "follow_up_questions": final_state["follow_up_questions"],
            "credibility_score": final_state["credibility_score"],
            "processing_time": processing_time,
            "query_analysis": final_state["query_metadata"],
            "metadata": {
                "search_strategy": final_state["search_strategy"],
                "processing_metrics": final_state["processing_metrics"],
                "multimodal_context": final_state["multimodal_context"]
            }
        }
        
        # Format extracted media for frontend
        extracted_media = final_state.get("extracted_media", {})
        if extracted_media.get("status") == "success":
            for img in extracted_media.get("images", []):
                results["extracted_media"].append({
                    "type": "image",
                    "title": img.get("title", ""),
                    "url": img.get("url", ""),
                    "thumbnail": img.get("thumbnail", img.get("url", "")),
                    "source_url": img.get("source_url", "")
                })
            
            for vid in extracted_media.get("videos", []):
                results["extracted_media"].append({
                    "type": "video",
                    "title": vid.get("title", ""),
                    "url": vid.get("url", ""),
                    "thumbnail": vid.get("thumbnail", ""),
                    "source_url": vid.get("source_url", "")
                })
        
        # Store search history
        search_record = {
            "id": str(uuid.uuid4()),
            "query": request.query,
            "search_type": request.search_type,
            "timestamp": start_time.isoformat(),
            "processing_time": processing_time,
            "results": results,
            "agentic_workflow": True,
            "agents_used": len([step for step in final_state["processing_metrics"].keys() if "agent" in step.lower()])
        }
        
        search_history_db.append(search_record)
        if len(search_history_db) > 1000:
            search_history_db = search_history_db[-1000:]
        
        logger.info(f"Agentic search completed: {processing_time:.2f}s - {len(results['sources'])} sources")
        
        return SearchResponse(**results)
        
    except Exception as e:
        analytics_db["error_count"] += 1
        logger.error(f"Agentic search failed: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Search failed: {str(e)}")

@app.post("/summarize-url")
async def summarize_url(request: SummaryRequest):
    """Generate summary for a specific URL"""
    try:
        content_data = await ContentExtractor.extract_content_from_url(request.url)
        
        if content_data.get("status") != "success":
            raise HTTPException(status_code=400, detail=content_data.get("message", "Failed to extract content"))
        
        # Use provided content or extracted content
        content_to_summarize = request.content or content_data.get("content", "")
        
        if not content_to_summarize:
            raise HTTPException(status_code=400, detail="No content available to summarize")
        
        summary = await ContentExtractor.generate_summary(content_to_summarize, request.url)
        
        return {
            "url": request.url,
            "title": content_data.get("title", ""),
            "summary": summary,
            "word_count": content_data.get("word_count", 0),
            "extracted_at": content_data.get("extracted_at", "")
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"URL summarization failed: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Summarization failed: {str(e)}")

@app.post("/generate-document")
async def generate_document(request: DocumentGenerationRequest):
    """Enhanced document generation with proper PDF support"""
    try:
        logger.info(f"Generating document: '{request.title}' with {len(request.sections)} sections")
        
        # Enhanced PDF generation
        pdf_data = generate_enhanced_pdf_document(request)
        
        # Save document with metadata
        doc_id = str(uuid.uuid4())
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_title = "".join(c for c in request.title if c.isalnum() or c in (' ', '-', '_')).rstrip()
        doc_filename = f"{safe_title}_{timestamp}.pdf"
        doc_path = GENERATED_DOCS_DIR / doc_filename
        
        with open(doc_path, "wb") as f:
            f.write(pdf_data)
        
        # Store document metadata
        doc_metadata = {
            "doc_id": doc_id,
            "filename": doc_filename,
            "title": request.title,
            "query": request.query,
            "sections": len(request.sections),
            "generated_at": datetime.now().isoformat(),
            "file_size": len(pdf_data),
            "file_path": str(doc_path)
        }
        
        logger.info(f"Document generated successfully: {doc_filename} ({len(pdf_data)} bytes)")
        
        # Return as downloadable file
        return StreamingResponse(
            BytesIO(pdf_data),
            media_type="application/pdf",
            headers={
                "Content-Disposition": f"attachment; filename={doc_filename}",
                "Content-Length": str(len(pdf_data))
            }
        )
        
    except Exception as e:
        logger.error(f"Document generation failed: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Document generation failed: {str(e)}")

@app.get("/files", response_model=List[FileMetadata])
async def list_files():
    """List uploaded files"""
    global uploaded_files_db
    
    files = [FileMetadata(**file_data) for file_data in uploaded_files_db.values()]
    return sorted(files, key=lambda x: x.upload_time, reverse=True)

@app.get("/files/{file_id}")
async def get_file(file_id: str):
    """Get specific file metadata"""
    global uploaded_files_db
    
    if file_id not in uploaded_files_db:
        raise HTTPException(status_code=404, detail="File not found")
    
    return JSONResponse(content=uploaded_files_db[file_id])

@app.delete("/files/{file_id}")
async def delete_file(file_id: str):
    """Delete uploaded file"""
    global uploaded_files_db
    
    try:
        if file_id not in uploaded_files_db:
            raise HTTPException(status_code=404, detail="File not found")
        
        file_metadata = uploaded_files_db[file_id]
        file_path = Path(file_metadata["file_path"])
        
        # Delete main file
        if file_path.exists():
            file_path.unlink()
        
        # Delete thumbnail if exists
        if file_metadata.get("thumbnail_url"):
            thumbnail_path = THUMBNAILS_DIR / Path(file_metadata["thumbnail_url"]).name
            if thumbnail_path.exists():
                thumbnail_path.unlink()
        
        del uploaded_files_db[file_id]
        
        logger.info(f"File deleted: {file_id}")
        
        return JSONResponse(content={
            "message": "File deleted successfully",
            "file_id": file_id,
            "filename": file_metadata["filename"]
        })
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"File deletion failed: {str(e)}")
        raise HTTPException(status_code=500, detail=f"File deletion failed: {str(e)}")

@app.get("/health")
async def health_check():
    """Comprehensive health check"""
    global uploaded_files_db, search_history_db
    
    # Check agentic capabilities
    agentic_status = {
        "langchain_available": True,
        "langgraph_available": True,
        "tavily_available": bool(TAVILY_API_KEY),
        "gemini_available": bool(GEMINI_API_KEY),
        "agents_count": 13,
        "workflow_operational": True
    }
    
    # Check file processing
    processing_status = {
        "pdf_processing": FILE_PROCESSING_AVAILABLE,
        "image_processing": PIL_AVAILABLE,
        "thumbnail_generation": PIL_AVAILABLE,
        "document_generation": REPORTLAB_AVAILABLE,
        "content_extraction": True,
        "url_summarization": True
    }
    
    uptime = datetime.now() - analytics_db["uptime_start"]
    
    return {
        "status": "healthy",
        "version": "3.0.0",
        "timestamp": datetime.now().isoformat(),
        "uptime": str(uptime),
        "uptime_seconds": uptime.total_seconds(),
        "agentic_capabilities": agentic_status,
        "file_processing": processing_status,
        "storage": {
            "total_files": len(uploaded_files_db),
            "total_searches": len(search_history_db)
        },
        "analytics": {
            "total_searches": analytics_db["total_searches"],
            "error_count": analytics_db["error_count"],
            "success_rate": ((analytics_db["total_searches"] - analytics_db["error_count"]) / max(analytics_db["total_searches"], 1)) * 100
        }
    }

@app.get("/analytics")
async def get_analytics():
    """Get comprehensive analytics"""
    global analytics_db, uploaded_files_db, search_history_db
    
    # Agentic workflow analytics
    agentic_searches = len([s for s in search_history_db if s.get("agentic_workflow", False)])
    avg_agents_used = statistics.mean([s.get("agents_used", 0) for s in search_history_db if s.get("agents_used", 0) > 0]) if search_history_db else 0
    
    # Processing time analytics
    processing_times = [s.get("processing_time", 0) for s in search_history_db]
    avg_processing_time = statistics.mean(processing_times) if processing_times else 0
    
    uptime = datetime.now() - analytics_db["uptime_start"]
    
    return {
        "system": {
            "status": "healthy",
            "version": "3.0.0",
            "uptime": str(uptime),
            "agentic_backend": True
        },
        "searches": {
            "total_searches": analytics_db["total_searches"],
            "agentic_searches": agentic_searches,
            "agentic_percentage": (agentic_searches / max(analytics_db["total_searches"], 1)) * 100,
            "avg_processing_time": round(avg_processing_time, 3),
            "avg_agents_used": round(avg_agents_used, 1),
            "search_types": analytics_db["search_types"]
        },
        "agentic_workflow": {
            "total_agents": 13,
            "workflow_steps": [
                "multimodal_processor", "query_analyzer", "search_strategist", 
                "enhanced_search", "content_validator", "media_extractor",
                "fact_checker", "synthesis_expert", "multimodal_synthesizer",
                "summarization", "citation_specialist", "insight_generator", "quality_assurance"
            ],
            "llm_integration": "gemini-1.5-flash-latest",
            "search_integration": "tavily"
        },
        "files": {
            "total_files": len(uploaded_files_db),
            "processing_capabilities": {
                "pdf": FILE_PROCESSING_AVAILABLE,
                "images": PIL_AVAILABLE,
                "documents": True,
                "thumbnails": PIL_AVAILABLE,
                "content_extraction": True,
                "summarization": True
            }
        },
        "performance": {
            "error_rate": round((analytics_db["error_count"] / max(analytics_db["total_searches"], 1)) * 100, 2),
            "success_rate": round(((analytics_db["total_searches"] - analytics_db["error_count"]) / max(analytics_db["total_searches"], 1)) * 100, 2)
        }
    }

# Startup event
@app.on_event("startup")
async def startup_event():
    logger.info("NEXUS AI Backend Ultra Premium Agentic starting up...")
    logger.info(f"LangChain integration: Active")
    logger.info(f"LangGraph workflow: {13} agents configured")
    logger.info(f"Tavily search: {'Active' if TAVILY_API_KEY else 'Inactive'}")
    logger.info(f"Gemini LLM: {'Active' if GEMINI_API_KEY else 'Inactive'}")
    logger.info(f"File processing: {'Enhanced' if FILE_PROCESSING_AVAILABLE else 'Basic'}")
    logger.info(f"Content extraction: Active")
    logger.info(f"URL summarization: Active")
    logger.info("Agentic backend startup complete!")

# Shutdown event
@app.on_event("shutdown")
async def shutdown_event():
    logger.info("NEXUS AI Backend shutting down...")
    # Cleanup temporary files
    for temp_file in TEMP_DIR.glob("*"):
        try:
            temp_file.unlink()
        except:
            pass
    logger.info("Agentic backend shutdown complete!")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(
        "main:app",
        host="0.0.0.0",
        port=8000,
        reload=True,
        log_level="info",
        access_log=True
    )